#!/usr/bin/env python3
"""Generate AG ONE Experion architecture Word document."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("AG ONE Experion", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph("Architecture, Components & Multi-Product Integration Guide")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    sub.runs[0].font.size = Pt(14)

    doc.add_paragraph()

    # --- 1. Executive Summary ---
    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "AG ONE Experion is a drop-in AI agent delivered as a single JavaScript SDK. "
        "It combines passive activity mining with an interactive recommendation and chatbot agent. "
        "Within the AG ONE product series, Experion is shared across six or more products "
        "(for example Learn, Work, and others). Each product team works independently on "
        "product-specific actions, knowledge bases, and response formats, while the Experion "
        "team owns only the Experion API and SDK."
    )
    doc.add_paragraph(
        "All user requests flow through AG ONE, which acts as the central orchestrator. "
        "AG ONE decides whether a request is an ACTION or a GENERATION, and routes execution "
        "to the correct product backend. Experion provides shared intelligence — intent routing, "
        "semantic cache, activity tracking, and proactive recommendations — without coupling "
        "product teams to Experion internals."
    )

    # --- 2. Design Principles ---
    add_heading(doc, "2. Architecture Design Principles", 1)
    principles = [
        ("Single SDK, many products", "One Experion JS SDK (hosted on Azure Blob Storage) is referenced by every AG ONE product via a single <script> tag. .NET web apps can also reference the same blob URL."),
        ("AG ONE as orchestrator", "Every chat, action, and generation request enters AG ONE first. AG ONE applies tenant/product context and routes downstream."),
        ("Team independence", "Product teams implement only their own Action Handlers and Generation/KB services. They integrate with AG ONE contracts, not Experion internals."),
        ("Experion API boundary", "The Experion team exposes /track and /process (plus SignalR). Other teams call AG ONE; AG ONE calls Experion where shared intelligence is needed."),
        ("Data-driven actions", "Action mappings live in SQL per tenant/product. New actions can be added without redeploying the SDK."),
        ("Cost-aware routing", "Semantic cache and pure-NLP action routing avoid unnecessary LLM calls."),
    ]
    for title_text, body in principles:
        add_bullet(doc, f" — {body}", bold_prefix=title_text)

    # --- 3. High-Level Architecture ---
    add_heading(doc, "3. High-Level Architecture", 1)
    doc.add_paragraph(
        "The architecture has four logical layers: Presentation (SDK), Platform (AG ONE), "
        "Shared Intelligence (Experion API), and Product Execution (per-product backends)."
    )

    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    headers = ["Layer", "Owner", "Responsibility"]
    rows_data = [
        ("Presentation", "Experion team", "JS SDK on Blob — tracking, chat UI, gestures, SignalR client"),
        ("Platform", "AG ONE platform team", "Gateway, tenant routing, ACTION vs GENERATION dispatch"),
        ("Shared intelligence", "Experion team", "NLP router, semantic cache, recommendations, audit"),
        ("Product execution", "Each product team", "Product-specific actions, KB/RAG, response shaping"),
    ]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "1F4E79")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row_data in enumerate(rows_data, start=1):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx].cells[c_idx].text = val

    doc.add_paragraph()

    # --- 4. Request Flow ---
    add_heading(doc, "4. Request Flow — AG ONE Orchestration", 1)

    add_heading(doc, "4.1 Chat / Action / Generation (/process)", 2)
    steps = [
        "User interacts via Experion SDK (chat message, circle gesture, or suggestion chip).",
        "SDK POSTs to AG ONE Gateway with tenant, product, userId, sessionId, and payload.",
        "AG ONE forwards to Experion /process for shared pipeline steps (embed, cache, intent classify).",
        "Experion returns intent: ACTION or GENERATION, plus matched actionKey or generation context.",
        "AG ONE routes execution:",
    ]
    for s in steps:
        add_bullet(doc, s)

    route_table = doc.add_table(rows=3, cols=3)
    route_table.style = "Table Grid"
    for i, h in enumerate(["Intent", "AG ONE routes to", "Product team delivers"]):
        cell = route_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "2E75B6")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    route_data = [
        ("ACTION", "Product N Action Handler API", "Execute add_to_cart, checkout, etc. via product APIs"),
        ("GENERATION", "Product N Generation / KB Service", "RAG answer using product KB, custom tone/format"),
    ]
    for r_idx, row_data in enumerate(route_data, start=1):
        for c_idx, val in enumerate(row_data):
            route_table.rows[r_idx].cells[c_idx].text = val

    doc.add_paragraph()
    doc.add_paragraph(
        "Experion persists conversation turns to Blob, updates SemanticCache on miss, "
        "and writes AuditLog. Product teams return responses in their contract format; "
        "AG ONE normalises the final payload for the SDK."
    )

    add_heading(doc, "4.2 Activity Tracking (/track)", 2)
    track_steps = [
        "SDK batches pageview, click, scroll, dwell, input, idle, and gesture events (~5s flush).",
        "SDK POSTs batch to AG ONE → Experion /track.",
        "Experion appends JSONL to activity blob and updates UserProfile in SQL.",
        "RecommendationEngine evaluates trigger rules and pushes nudges via SignalR.",
    ]
    for s in track_steps:
        add_bullet(doc, s)

    # --- 5. Core Components ---
    add_heading(doc, "5. Core Components", 1)

    add_heading(doc, "5.1 Frontend — Experion JS SDK (Blob-hosted)", 2)
    sdk_items = [
        "Single <script> tag: data-tenant, data-product, data-api (AG ONE URL), data-user-id",
        "Identity Resolver — logged-in UserID or anon-<id> + IP/region cohort",
        "Activity Tracker — passive behavioural capture",
        "Circle Gesture Capture — Alt+drag region → DOM text extraction",
        "Chat Sidebar — avatar, panel, suggestion chips",
        "Proactive Nudge Toast — SignalR-driven recommendations",
        "sendBeacon flush on page unload",
        "Referenced by all AG ONE product frontends (.NET MVC/Razor, SPA, static sites)",
    ]
    for item in sdk_items:
        add_bullet(doc, item)

    add_heading(doc, "5.2 Platform — AG ONE Gateway", 2)
    agone_items = [
        "Single entry point for all AG ONE products",
        "Tenant + product resolution (Learn, Work, Product 3…)",
        "Auth / API key validation per product",
        "Orchestration: call Experion, then dispatch to product backends",
        "Response aggregation and SDK-compatible envelope",
        "Rate limiting and circuit breaking per product",
    ]
    for item in agone_items:
        add_bullet(doc, item)

    add_heading(doc, "5.3 Backend — Experion API (.NET 8)", 2)
    api_items = [
        "POST /track — ingest activity events",
        "POST /process — embed, cache, intent route, persist",
        "ExperionService — flat pipeline, one method per step",
        "RecommendationEngine — activity → LLM nudge → SignalR",
        "ActionWorker — consumes action queue (Service Bus in prod); AG ONE may also enqueue product jobs",
        "SignalR hub — nudges and action acknowledgements",
    ]
    for item in api_items:
        add_bullet(doc, item)

    add_heading(doc, "5.4 Product Backends (per team)", 2)
    doc.add_paragraph(
        "Each AG ONE product team owns two integration surfaces registered with AG ONE:"
    )
    add_bullet(doc, "Action Handler — executes data-driven actions for that product's domain", bold_prefix="")
    add_bullet(doc, "Generation Service — KB retrieval + LLM with product-specific prompts and response schema", bold_prefix="")
    doc.add_paragraph(
        "Teams never modify Experion code. They publish OpenAPI contracts to AG ONE and "
        "implement handlers behind their own APIs."
    )

    # --- 6. Storage ---
    add_heading(doc, "6. Storage", 1)
    storage_table = doc.add_table(rows=6, cols=3)
    storage_table.style = "Table Grid"
    for i, h in enumerate(["Store", "Contents", "Used by"]):
        cell = storage_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "1F4E79")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    storage_data = [
        ("Blob (JSONL)", "conversations/, activity/", "Experion"),
        ("SQL Server", "ActionMappings, SemanticCache, UserProfile, AuditLog, TenantConfig", "Experion + AG ONE"),
        ("Azure AI Search", "KB index, semantic-cache vectors", "Experion + product KB indexes"),
        ("Azure OpenAI", "Embeddings + GPT-4.1", "Experion + product generation"),
        ("Service Bus", "Action job queue", "Experion ActionWorker + product workers"),
    ]
    for r_idx, row_data in enumerate(storage_data, start=1):
        for c_idx, val in enumerate(row_data):
            storage_table.rows[r_idx].cells[c_idx].text = val

    # --- 7. /process pipeline detail ---
    add_heading(doc, "7. Experion /process Pipeline (Shared Intelligence)", 1)
    pipeline = [
        ("Step 1 · Embed Query", "Turn user text (and optional capturedText from circle gesture) into a vector."),
        ("Step 2 · Semantic Cache", "Cosine search over SemanticCache; HIT returns cached answer (< 15 ms)."),
        ("Step 3 · NLP Intent Router", "Hybrid score (embedding + Jaccard) vs ActionMappings → ACTION or GENERATION."),
        ("Step 4a · Action path", "Return actionKey + metadata to AG ONE; AG ONE dispatches to product Action Handler."),
        ("Step 4b · Generation path", "Return generation intent to AG ONE; product KB service runs RAG (or Experion RAG for shared KB)."),
        ("Step 5 · Persist", "JSONL conversation lines to Blob; cache write on miss; AuditLog row."),
    ]
    for title_text, body in pipeline:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(title_text)
        run.bold = True
        p.add_run(f" — {body}")

    # --- 8. Multi-product matrix ---
    add_heading(doc, "8. Multi-Product Routing Matrix (Example)", 1)
    doc.add_paragraph(
        "AG ONE maintains a routing table per product. Action and generation responses "
        "may differ per product (JSON schema, UI chips, redirect URLs)."
    )
    matrix = doc.add_table(rows=4, cols=5)
    matrix.style = "Table Grid"
    matrix_headers = ["Product", "Tenant ID", "Action endpoint", "Generation endpoint", "Notes"]
    for i, h in enumerate(matrix_headers):
        cell = matrix.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "2E75B6")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    matrix_data = [
        ("AG ONE Learn", "learn", "learn-api/actions", "learn-api/generate", "Course/cart actions"),
        ("AG ONE Work", "work", "work-api/actions", "work-api/generate", "Task/project actions"),
        ("Product N", "product-n", "product-n-api/actions", "product-n-api/generate", "Team-owned contract"),
    ]
    for r_idx, row_data in enumerate(matrix_data, start=1):
        for c_idx, val in enumerate(row_data):
            matrix.rows[r_idx].cells[c_idx].text = val

    # --- 9. Actions ---
    add_heading(doc, "9. Kinds of Actions (Data-Driven)", 1)
    doc.add_paragraph(
        "ActionMappings in SQL are scoped per tenant/product. Seeded examples:"
    )
    actions = [
        "add_to_cart — add viewed product to cart",
        "checkout — navigate to payment flow",
        "apply_coupon — apply discount code",
        "open_profile — account/settings",
        "contact_support — ticket or email",
        "logout — sign out",
    ]
    for a in actions:
        add_bullet(doc, a)
    doc.add_paragraph(
        "Each mapping includes: actionKey, description, phrases, embedding, targetEndpoint, enabled. "
        "Product teams register targetEndpoint with AG ONE; Experion only classifies intent."
    )

    # --- 10. Generation ---
    add_heading(doc, "10. Kinds of Questions (Generation Path)", 1)
    gen_items = [
        "Product overviews, pricing, plans (product KB)",
        "Policies — refund, return, cancellation",
        "How-tos — cart, checkout, billing, account",
        "Follow-ups using conversation history from Blob",
        "Circle-gesture 'what is this?' with capturedText DOM fragment",
    ]
    for g in gen_items:
        add_bullet(doc, g)

    # --- 11. Recommendations ---
    add_heading(doc, "11. Proactive Recommendations", 1)
    doc.add_paragraph(
        "Behaviour-triggered nudges (not user-initiated), throttled per user:"
    )
    nudges = [
        "You've viewed 3 products — compare them?",
        "Stuck on a form — walkthrough?",
        "Cart idle — ready to checkout?",
    ]
    for n in nudges:
        add_bullet(doc, n)

    # --- 12. Identity ---
    add_heading(doc, "12. Identity Model", 1)
    identity = [
        "Logged-in → real UserID from host application",
        "Anonymous → anon-<id> in localStorage + IP region cohort",
        "SessionID = <userId>__<random> from /identify",
        "Session ties activity events and conversation JSONL",
    ]
    for i in identity:
        add_bullet(doc, i)

    # --- 13. Team boundaries ---
    add_heading(doc, "13. Team Boundaries & Integration Contracts", 1)

    teams_table = doc.add_table(rows=4, cols=3)
    teams_table.style = "Table Grid"
    for i, h in enumerate(["Team", "Owns", "Integrates via"]):
        cell = teams_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "1F4E79")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    teams_data = [
        ("Experion", "JS SDK, Experion API, SignalR, shared NLP/cache", "Blob script URL; /track, /process"),
        ("AG ONE Platform", "Gateway, routing table, auth, orchestration", "Product registration API"),
        ("Product team (e.g. Learn)", "Action Handler, Generation Service, product KB", "AG ONE product contracts only"),
    ]
    for r_idx, row_data in enumerate(teams_data, start=1):
        for c_idx, val in enumerate(row_data):
            teams_table.rows[r_idx].cells[c_idx].text = val

    # --- 14. Optimisations ---
    add_heading(doc, "14. Cost & Latency Optimisations", 1)
    opts = [
        "Semantic cache first — repeated paraphrases skip LLM (HIT ~< 15 ms vs 60–2000 ms)",
        "Action path — no LLM; NLP → SQL → queue → product handler",
        "/track fully async — never blocks chat",
        "Conversation history in Blob — SQL stays small and indexed",
    ]
    for o in opts:
        add_bullet(doc, o)

    # --- 15. Production ---
    add_heading(doc, "15. Production Swap-In Summary", 1)
    prod = [
        "Mock LLM → Azure OpenAI",
        "Local JSONL files → Azure Blob Append Blobs",
        "SQLite → SQL Server",
        "In-process queue → Azure Service Bus",
        "In-process SignalR → Azure SignalR Service",
        "Naive KB → Azure AI Search hybrid retrieval",
    ]
    for p in prod:
        add_bullet(doc, p)

    # --- 16. Onboarding ---
    add_heading(doc, "16. Tenant & Product Onboarding", 1)
    onboard = [
        "Experion: Add TenantConfig + ActionMappings for the product",
        "Product team: Register action + generation endpoints with AG ONE",
        "Product team: Add <script> tag with data-tenant and data-product",
        "AG ONE: Add routing row linking tenant → product backends",
        "Verify /track, /process, and SignalR end-to-end",
    ]
    for i, step in enumerate(onboard, start=1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)

  # --- 17. Diagram reference ---
    add_heading(doc, "17. Architecture Diagram", 1)
    doc.add_paragraph(
        "See the companion file ag-one-experion-architecture.drawio in the docs/ag-one-experion folder. "
        "Open with draw.io (diagrams.net) or the Draw.io desktop app. "
        "The diagram shows SDK → AG ONE → Experion → Product backends, team boundaries, and data stores."
    )

    doc.add_paragraph()
    footer = doc.add_paragraph("Document generated for AG ONE Experion — multi-product architecture.")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].italic = True
    footer.runs[0].font.size = Pt(9)

    return doc


if __name__ == "__main__":
    doc = build_document()
    out = "docs/ag-one-experion/AG_ONE_Experion_Architecture.docx"
    doc.save(out)
    print(f"Wrote {out}")
