#!/usr/bin/env python3
"""Generate AG ONE Marketplace Architecture Word doc and draw.io file."""

import html
import uuid
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(__file__).parent
DOCX_PATH = OUT_DIR / "AG-ONE-Marketplace-Architecture.docx"
DRAWIO_PATH = OUT_DIR / "AG-ONE-Marketplace-Architecture.drawio"


def set_cell_shading(cell, color_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "3B82F6")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    return table


def build_word_document():
    doc = Document()
    title = doc.add_heading("AG ONE Series — Azure Marketplace SaaS Integration", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Production Architecture & Design Document")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(30, 58, 95)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version 1.0 | Generated {datetime.utcnow().strftime('%Y-%m-%d')} | AG ONE Gateway + AG ONE Learn (SaaS)").italic = True

    doc.add_paragraph()
    doc.add_paragraph(
        "This document describes the production-ready architecture for publishing an AG ONE Series product "
        "(AG ONE Learn) on Azure Marketplace while maintaining AG ONE as the central authentication gateway."
    )

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "AG ONE acts as the central identity and entitlement gateway for all products (Work, Learn, Safe). "
        "Azure Marketplace handles billing and subscription lifecycle. AG ONE resolves marketplace tokens, "
        "activates subscriptions via the SaaS Fulfillment API, provisions tenants in core.*, and issues "
        "tenant-scoped JWTs shared across products via the agone_sso_token cookie."
    )
    add_bullet(doc, "Single Azure SQL database with schema separation: core.*, learn.*, work.*, safe.*")
    add_bullet(doc, "JWT authentication via Microsoft Entra ID + AgOneSsoMiddleware")
    add_bullet(doc, "Marketplace subscription maps 1:1 to AG ONE tenant for AG ONE Learn")
    add_bullet(doc, "Companion draw.io file: AG-ONE-Marketplace-Architecture.drawio (9 diagram pages)")

    # 2. High-Level Architecture
    add_heading(doc, "2. High-Level Architecture", 1)
    doc.add_paragraph("See draw.io page: 1-High-Level-Architecture")
    add_table(doc, ["Layer", "Component", "Responsibility"], [
        ["External", "Azure Marketplace", "Offer catalog, billing, subscription lifecycle, redirect to landing page"],
        ["Gateway", "AG ONE Gateway", "Landing page, token resolve, webhooks, tenant provisioning, JWT issuance, central logout"],
        ["Product", "AG ONE Learn", "SaaS product UI/API; validates AG ONE JWT; tenant-scoped learn.* data"],
        ["Product", "AG ONE Work / Safe", "Unchanged; SSO via AG ONE; not exposed via Marketplace unless separately listed"],
        ["Data", "Azure SQL", "Single DB; marketplace metadata in core.*; product data in learn.*"],
        ["Events", "Azure Service Bus", "Async webhook processing and tenant provisioning with retries"],
        ["Secrets", "Azure Key Vault", "Fulfillment API credentials, connection strings"],
    ])

    # 3. End-to-End Flow
    add_heading(doc, "3. End-to-End Flow", 1)
    doc.add_paragraph("See draw.io page: 2-End-to-End-Sequence")
    steps = [
        "Customer subscribes to AG ONE Learn on Azure Marketplace.",
        "Azure creates SaaS subscription (PendingFulfillmentStart) and redirects to AG ONE landing page with token.",
        "AG ONE resolves token via Fulfillment API POST /api/saas/subscriptions/resolve.",
        "AG ONE activates subscription via POST /api/saas/subscriptions/{id}/activate.",
        "TenantProvisioningWorker creates core.Tenants, core.Subscriptions, admin user, and roles.",
        "MarketplaceSubscription.AzureSubscriptionId maps to Subscription → Tenant → Product (Learn).",
        "User logs in via AG ONE (Entra ID OIDC).",
        "AG ONE issues JWT in agone_sso_token cookie with tenant_id and entitlements.",
        "User launches Learn from My Products; AgOneSsoMiddleware validates JWT.",
        "Learn queries learn.* schema filtered by TenantId.",
    ]
    for i, s in enumerate(steps, 1):
        add_bullet(doc, f"{i}. {s}")

    # 4. SaaS Onboarding Flow
    add_heading(doc, "4. SaaS Onboarding Flow", 1)
    doc.add_paragraph("See draw.io page: 3-SaaS-Onboarding-Flow")
    add_table(doc, ["Endpoint", "Auth", "Purpose"], [
        ["GET /marketplace/landing", "Anonymous", "Receives Marketplace redirect token"],
        ["POST /api/marketplace/webhook", "Marketplace signature", "Receives lifecycle webhooks"],
        ["GET /api/marketplace/subscriptions/{id}", "Admin", "Internal subscription status"],
        ["POST /api/marketplace/onboarding/complete", "Authenticated", "Post-login onboarding wizard"],
    ])
    doc.add_paragraph("Token Exchange (Fulfillment API v2):")
    doc.add_paragraph(
        "POST https://marketplaceapi.microsoft.com/api/saas/subscriptions/resolve?api-version=2018-08-31\n"
        "Authorization: Bearer {AAD_token_for_Fulfillment_API}\n"
        'Body: { "token": "{marketplaceToken_from_redirect}" }',
        style="Intense Quote",
    )
    doc.add_paragraph(
        "Persist: id (Azure subscription ID), offerId, planId, quantity, purchaser.email, "
        "beneficiary.email, saasSubscriptionStatus."
    )

    # 5. Database Design
    add_heading(doc, "5. Database Design Changes", 1)
    doc.add_paragraph("See draw.io page: 4-Database-ER-Diagram")
    add_heading(doc, "5.1 Extend core.Tenants", 2)
    add_table(doc, ["Column", "Type", "Description"], [
        ["Slug", "NVARCHAR(100)", "URL-friendly tenant identifier (unique)"],
        ["ProvisioningSource", "NVARCHAR(50)", "Direct | Marketplace | Partner"],
        ["Status", "NVARCHAR(50)", "Provisioning | Active | Suspended | Cancelled"],
        ["AzurePurchaserTenantId", "NVARCHAR(100)", "Azure AD tenant of purchaser"],
        ["OnboardingCompletedAt", "DATETIME2", "When onboarding wizard completed"],
    ])

    add_heading(doc, "5.2 Extend core.Products", 2)
    add_table(doc, ["Column", "Type", "Description"], [
        ["MarketplaceOfferId", "NVARCHAR(200)", "e.g. agonelearn"],
        ["MarketplacePublisherId", "NVARCHAR(200)", "Partner Center publisher ID"],
        ["LandingPagePath", "NVARCHAR(500)", "/marketplace/landing"],
        ["IsMarketplaceEnabled", "BIT", "Whether product is on Marketplace"],
    ])

    add_heading(doc, "5.3 Extend core.Subscriptions", 2)
    add_table(doc, ["Column", "Type", "Description"], [
        ["BillingSource", "NVARCHAR(50)", "Direct | AzureMarketplace"],
        ["ExternalSubscriptionId", "NVARCHAR(200)", "Azure subscription ID"],
        ["SuspendedAt", "DATETIME2", "When suspended"],
        ["CancelledAt", "DATETIME2", "When cancelled"],
    ])

    add_heading(doc, "5.4 New Table: core.MarketplaceSubscriptions", 2)
    add_table(doc, ["Column", "Type", "Description"], [
        ["Id", "UNIQUEIDENTIFIER PK", "Internal ID"],
        ["SubscriptionId", "UNIQUEIDENTIFIER FK", "Links to core.Subscriptions"],
        ["AzureSubscriptionId", "NVARCHAR(200) UK", "Marketplace subscription ID"],
        ["OfferId", "NVARCHAR(200)", "Marketplace offer ID"],
        ["PlanId", "NVARCHAR(100)", "Selected plan"],
        ["ProductId", "UNIQUEIDENTIFIER FK", "AG ONE Learn product ID"],
        ["SaasStatus", "NVARCHAR(50)", "PendingFulfillmentStart | Subscribed | Suspended | Unsubscribed"],
        ["PurchaserEmail", "NVARCHAR(256)", "Buyer email"],
        ["BeneficiaryEmail", "NVARCHAR(256)", "End-user email if different"],
        ["Quantity", "INT", "License seats"],
        ["ResolvedAt", "DATETIME2", "Token resolve timestamp"],
        ["ActivatedAt", "DATETIME2", "Activation timestamp"],
    ])

    add_heading(doc, "5.5 New Table: core.MarketplaceWebhookEvents", 2)
    add_table(doc, ["Column", "Type", "Description"], [
        ["Id", "UNIQUEIDENTIFIER PK", "Event ID"],
        ["AzureSubscriptionId", "NVARCHAR(200)", "Related subscription"],
        ["Action", "NVARCHAR(100)", "Unsubscribe | ChangePlan | ChangeQuantity | Suspend | Reinstate"],
        ["Status", "NVARCHAR(50)", "Received | Processed | Failed"],
        ["IdempotencyKey", "NVARCHAR(200) UK", "Prevents duplicate processing"],
        ["PayloadJson", "NVARCHAR(MAX)", "Raw webhook payload"],
    ])

    add_heading(doc, "5.6 New Table: core.TenantProvisioningJobs", 2)
    add_table(doc, ["Column", "Type", "Description"], [
        ["Id", "UNIQUEIDENTIFIER PK", "Job ID"],
        ["MarketplaceSubscriptionId", "UNIQUEIDENTIFIER FK", "Source subscription"],
        ["TenantId", "UNIQUEIDENTIFIER", "Created tenant (after success)"],
        ["Status", "NVARCHAR(50)", "Queued | Running | Completed | Failed"],
        ["Step", "NVARCHAR(100)", "Current provisioning step"],
        ["AttemptCount", "INT", "Retry count"],
    ])

    add_heading(doc, "5.7 Mapping Chain", 2)
    doc.add_paragraph(
        "Azure Marketplace subscriptionId → core.MarketplaceSubscriptions.AzureSubscriptionId → "
        "core.MarketplaceSubscriptions.SubscriptionId → core.Subscriptions (TenantId + ProductId) → "
        "core.Tenants.Id → learn.* (TenantId on all rows)"
    )

    # 6. Integration Design
    add_heading(doc, "6. Azure Marketplace Integration Design", 1)
    doc.add_paragraph("See draw.io pages: 5-Integration-Design, 9-Subscription-Lifecycle")
    add_table(doc, ["Fulfillment API Method", "Endpoint", "When Used"], [
        ["ResolveToken", "POST /subscriptions/resolve", "Landing page token exchange"],
        ["Activate", "POST /subscriptions/{id}/activate", "After successful resolve"],
        ["GetSubscription", "GET /subscriptions/{id}", "Reconciliation, webhook validation"],
        ["Update", "PATCH /subscriptions/{id}", "Plan/quantity change webhook"],
    ])
    add_heading(doc, "6.1 Webhook Lifecycle Actions", 2)
    add_table(doc, ["Webhook Action", "AG ONE Behavior"], [
        ["Subscribe", "Handled on landing page (resolve + activate); webhook as backup"],
        ["ChangePlan", "Update PlanName; call PATCH Fulfillment API; adjust feature flags"],
        ["ChangeQuantity", "Update MaxLicenses; enforce seat limits on user invites"],
        ["Suspend", "Set Subscription + Tenant Status = Suspended; block product login"],
        ["Reinstate", "Restore Status = Active"],
        ["Unsubscribe", "Set Cancelled; soft-disable access; retain data per retention policy"],
    ])
    doc.add_paragraph(
        "Webhook URL (Partner Center): https://agone.example.com/api/marketplace/webhook\n"
        "Return HTTP 200 quickly; process heavy work via Azure Service Bus."
    )

    # 7. Authentication Flow
    add_heading(doc, "7. Authentication Flow Update", 1)
    doc.add_paragraph("See draw.io page: 6-Authentication-Flow")
    add_heading(doc, "7.1 Extended JWT Claims", 2)
    add_table(doc, ["Claim", "Example", "Purpose"], [
        ["tenant_id", "guid", "Row-level isolation (existing)"],
        ["tenant_slug", "acme-corp", "Friendly URLs"],
        ["product_id", "Learn GUID", "Active product context"],
        ["subscription_id", "guid", "Internal subscription reference"],
        ["subscription_status", "Active / Suspended", "Product-level gate"],
        ["billing_source", "AzureMarketplace", "UI/billing behavior"],
        ["marketplace_subscription_id", "Azure string", "Support/billing correlation"],
        ["max_licenses", "25", "Seat enforcement"],
        ["permission", "learn.path.read", "RBAC (existing)"],
    ])
    add_heading(doc, "7.2 Multi-Tenant Rules", 2)
    rules = [
        "One Marketplace subscription = one AG ONE tenant (default isolation model).",
        "Purchaser email becomes initial TenantOwner in core.Users.",
        "Beneficiary email used when different from purchaser.",
        "Before JWT issuance: verify Tenant.Status=Active, Subscription.Status=Active, Marketplace SaasStatus=Subscribed.",
        "Central logout unchanged: any product logout redirects to AG ONE; cookie cleared domain-wide.",
        "Seat limits: UsedLicenses < MaxLicenses on user invite.",
    ]
    for r in rules:
        add_bullet(doc, r)

    # 8. System Flow
    add_heading(doc, "8. System Flow (Step-by-Step)", 1)
    doc.add_paragraph("See draw.io page: 7-System-Flow-Steps")
    flow = [
        ("User browses Azure Marketplace", "Customer discovers AG ONE Learn offer"),
        ("Subscribe", "Azure creates PendingFulfillmentStart subscription"),
        ("Redirect", "Azure sends user to AG ONE landing page with token"),
        ("Resolve", "AG ONE calls Fulfillment API /resolve"),
        ("Activate", "AG ONE calls /activate within 30-day window"),
        ("Provision Tenant", "Create core.Tenants with ProvisioningSource=Marketplace"),
        ("Create Subscription", "core.Subscriptions for AGOneLearn, Status=Active"),
        ("Map", "AzureSubscriptionId → TenantId → ProductId"),
        ("Create Admin", "User + TenantOwner + LearnProdAdmin roles"),
        ("Seed learn.*", "Default tenant data in learn schema"),
        ("Login", "User authenticates via AG ONE / Entra ID"),
        ("JWT", "agone_sso_token cookie with tenant_id"),
        ("Launch Learn", "From My Products page"),
        ("Validate SSO", "AgOneSsoMiddleware on Learn app"),
        ("Access", "Tenant-scoped learn.* queries"),
    ]
    add_table(doc, ["Step", "Description"], flow)

    # 9. Deployment
    add_heading(doc, "9. Deployment Architecture", 1)
    doc.add_paragraph("See draw.io page: 8-Deployment-Architecture")
    add_table(doc, ["Component", "Azure Service", "Notes"], [
        ["AG ONE Gateway", "Azure App Service (Premium v3)", "Landing page + webhooks; HA required"],
        ["AG ONE Learn", "App Service or AKS", "Start App Service; scale to AKS"],
        ["Database", "Azure SQL Business Critical", "TDE, AAD auth, schema isolation"],
        ["Secrets", "Key Vault", "Fulfillment API creds via managed identity"],
        ["Async Events", "Service Bus", "marketplace-events queue/topic"],
        ["Monitoring", "Application Insights", "Webhook failures, provisioning DLQ alerts"],
        ["CDN (optional)", "Azure Front Door", "Global Marketplace purchasers"],
    ])
    add_heading(doc, "9.1 New API Structure", 2)
    apis = [
        "AGOne.Api/Controllers/MarketplaceLandingController.cs — GET /marketplace/landing",
        "AGOne.Api/Controllers/MarketplaceWebhookController.cs — POST /api/marketplace/webhook",
        "AGOne.Infrastructure/Services/MarketplaceFulfillmentService.cs",
        "AGOne.Infrastructure/Services/MarketplaceLifecycleService.cs",
        "AGOne.Infrastructure/Services/TenantProvisioningService.cs",
        "AGOne.Infrastructure/Workers/TenantProvisioningWorker.cs — Service Bus triggered",
    ]
    for a in apis:
        add_bullet(doc, a)

    # 10. Design Considerations
    add_heading(doc, "10. Design Considerations", 1)
    add_heading(doc, "10.1 Multi-Tenancy", 2)
    add_table(doc, ["Concern", "Approach"], [
        ["Data isolation", "TenantId on every learn.* row; EF global query filters"],
        ["Schema separation", "learn.* separate from work.*/safe.*; Marketplace provisions Learn only"],
        ["Tenant boundary", "1 Azure subscription = 1 AG ONE tenant"],
        ["Noisy neighbor", "Rate limiting per tenant at API Management / middleware"],
    ])
    add_heading(doc, "10.2 Security", 2)
    add_table(doc, ["Concern", "Approach"], [
        ["Webhook authenticity", "Validate Marketplace JWT/signature; idempotency keys"],
        ["Token handling", "Marketplace token single-use; never log full token"],
        ["Fulfillment API creds", "Key Vault + managed identity; rotate secrets"],
        ["JWT", "Short-lived; HttpOnly Secure SameSite=None cookie"],
        ["Tenant escalation", "Never trust client-supplied tenant_id; always from validated JWT"],
    ])
    add_heading(doc, "10.3 Scalability & Operations", 2)
    add_table(doc, ["Concern", "Approach"], [
        ["Provisioning burst", "Async Service Bus workers; horizontal scale"],
        ["Webhook spikes", "Queue-first; return 200 immediately"],
        ["Reconciliation", "Nightly Fulfillment API GET vs DB sync job"],
        ["DLQ monitoring", "Alert on failed provisioning jobs"],
        ["Environments", "Separate Landing Page + Webhook URLs per dev/uat/prod in Partner Center"],
    ])

    # 11. Implementation Phases
    add_heading(doc, "11. Implementation Phases", 1)
    add_table(doc, ["Phase", "Deliverables"], [
        ["Phase 1 — Foundation", "DB migrations, MarketplaceFulfillmentService, landing page, activate flow"],
        ["Phase 2 — Provisioning", "Tenant/user/role auto-provision, learn.* seed, welcome email"],
        ["Phase 3 — Lifecycle", "Webhook handler, suspend/cancel/reinstate, seat management"],
        ["Phase 4 — Auth Hardening", "Extended JWT claims, subscription gate on login, Learn middleware update"],
        ["Phase 5 — Operations", "Reconciliation, dashboards, Partner Center go-live"],
    ])

    doc.add_page_break()
    add_heading(doc, "Appendix A — SQL DDL (core.MarketplaceSubscriptions)", 1)
    ddl = """
CREATE TABLE [core].[MarketplaceSubscriptions] (
    [Id]                    UNIQUEIDENTIFIER NOT NULL DEFAULT NEWSEQUENTIALID(),
    [SubscriptionId]        UNIQUEIDENTIFIER NULL,
    [AzureSubscriptionId]   NVARCHAR(200)    NOT NULL,
    [OfferId]               NVARCHAR(200)    NOT NULL,
    [PlanId]                NVARCHAR(100)    NOT NULL,
    [ProductId]             UNIQUEIDENTIFIER NOT NULL,
    [SaasStatus]            NVARCHAR(50)     NOT NULL,
    [PurchaserEmail]        NVARCHAR(256)    NOT NULL,
    [BeneficiaryEmail]      NVARCHAR(256)    NULL,
    [PurchaserTenantId]     NVARCHAR(100)    NULL,
    [Quantity]              INT              NOT NULL DEFAULT 1,
    [ResolvedPayloadJson]   NVARCHAR(MAX)    NULL,
    [ResolvedAt]            DATETIME2        NOT NULL,
    [ActivatedAt]           DATETIME2        NULL,
    [LastWebhookAt]         DATETIME2        NULL,
    [CreatedAt]             DATETIME2        NOT NULL DEFAULT GETUTCDATE(),
    [UpdatedAt]             DATETIME2        NULL,
    [IsDeleted]             BIT              NOT NULL DEFAULT 0,
    CONSTRAINT [PK_MarketplaceSubscriptions] PRIMARY KEY ([Id]),
    CONSTRAINT [FK_MarketplaceSubscriptions_Products] FOREIGN KEY ([ProductId]) REFERENCES [core].[Products]([Id]),
    CONSTRAINT [FK_MarketplaceSubscriptions_Subscriptions] FOREIGN KEY ([SubscriptionId]) REFERENCES [core].[Subscriptions]([Id]),
    CONSTRAINT [UQ_MarketplaceSubscriptions_AzureSubscriptionId] UNIQUE ([AzureSubscriptionId])
);
"""
    doc.add_paragraph(ddl.strip(), style="Intense Quote")

    add_heading(doc, "Appendix B — Draw.io Diagram Index", 1)
    pages = [
        "1-High-Level-Architecture",
        "2-End-to-End-Sequence",
        "3-SaaS-Onboarding-Flow",
        "4-Database-ER-Diagram",
        "5-Integration-Design",
        "6-Authentication-Flow",
        "7-System-Flow-Steps",
        "8-Deployment-Architecture",
        "9-Subscription-Lifecycle",
    ]
    for p in pages:
        add_bullet(doc, p)

    doc.save(DOCX_PATH)
    print(f"Created: {DOCX_PATH}")


# ─── draw.io helpers ─────────────────────────────────────────────────────────

def esc(text: str) -> str:
    return html.escape(str(text), quote=True)


def rect(cid, x, y, w, h, label, fill="#dae8fc", stroke="#6c8ebf", parent="1", rounded=1):
    style = (
        f"rounded={1 if rounded else 0};whiteSpace=wrap;html=1;"
        f"fillColor={fill};strokeColor={stroke};fontSize=12;fontStyle=0"
    )
    return (
        f'<mxCell id="{cid}" value="{esc(label)}" style="{style}" vertex="1" parent="{parent}">'
        f'<mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/></mxCell>'
    )


def cylinder(cid, x, y, w, h, label, fill="#fff2cc", stroke="#d6b656", parent="1"):
    style = (
        "shape=cylinder3;whiteSpace=wrap;html=1;boundedLbl=1;backgroundOutline=1;"
        f"size=15;fillColor={fill};strokeColor={stroke};fontSize=12"
    )
    return (
        f'<mxCell id="{cid}" value="{esc(label)}" style="{style}" vertex="1" parent="{parent}">'
        f'<mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/></mxCell>'
    )


def cloud(cid, x, y, w, h, label, fill="#f5f5f5", stroke="#666666", parent="1"):
    style = f"ellipse;shape=cloud;whiteSpace=wrap;html=1;fillColor={fill};strokeColor={stroke};fontSize=12"
    return (
        f'<mxCell id="{cid}" value="{esc(label)}" style="{style}" vertex="1" parent="{parent}">'
        f'<mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/></mxCell>'
    )


def swimlane(cid, x, y, w, h, label, fill="#e1d5e7", stroke="#9673a6", parent="1"):
    style = (
        "swimlane;whiteSpace=wrap;html=1;startSize=30;"
        f"fillColor={fill};strokeColor={stroke};fontSize=13;fontStyle=1"
    )
    return (
        f'<mxCell id="{cid}" value="{esc(label)}" style="{style}" vertex="1" parent="{parent}">'
        f'<mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/></mxCell>'
    )


def edge(eid, source, target, label="", dashed=False, parent="1"):
    style = "edgeStyle=orthogonalEdgeStyle;rounded=1;orthogonalLoop=1;jettySize=auto;html=1;fontSize=11"
    if dashed:
        style += ";dashed=1"
    lbl = f' value="{esc(label)}"' if label else ""
    return (
        f'<mxCell id="{eid}" style="{style}" edge="1" parent="{parent}" source="{source}" target="{target}"{lbl}>'
        f'<mxGeometry relative="1" as="geometry"/></mxCell>'
    )


def diamond(cid, x, y, w, h, label, parent="1"):
    style = "rhombus;whiteSpace=wrap;html=1;fillColor=#fff2cc;strokeColor=#d6b656;fontSize=11"
    return (
        f'<mxCell id="{cid}" value="{esc(label)}" style="{style}" vertex="1" parent="{parent}">'
        f'<mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/></mxCell>'
    )


def page1_high_level():
    cells = []
    cells.append(swimlane("z1", 20, 20, 280, 520, "External", "#f8cecc", "#b85450"))
    cells.append(cloud("u1", 60, 60, 200, 80, "Customer / Purchaser", parent="z1"))
    cells.append(rect("u2", 40, 180, 220, 70, "Azure Marketplace\n(Partner Center Offer)", "#f8cecc", "#b85450", "z1"))
    cells.append(rect("u3", 40, 300, 220, 70, "Microsoft Entra ID", "#f8cecc", "#b85450", "z1"))

    cells.append(swimlane("z2", 320, 20, 820, 520, "Azure Subscription (Publisher)", "#d5e8d4", "#82b366"))
    cells.append(swimlane("z3", 340, 60, 780, 120, "Edge Layer", "#fff2cc", "#d6b656", "z2"))
    cells.append(rect("e1", 380, 100, 160, 50, "App Gateway / WAF", parent="z3"))
    cells.append(rect("e2", 580, 100, 160, 50, "API Management", parent="z3"))

    cells.append(swimlane("z4", 340, 200, 780, 180, "Compute", "#dae8fc", "#6c8ebf", "z2"))
    cells.append(rect("c1", 360, 250, 200, 90, "AG ONE Gateway\nAuth · IAM · Marketplace\nLanding + Webhooks", "#3b82f6", "#1e3a5f", "z4"))
    cells.append(rect("c2", 590, 250, 150, 90, "AG ONE Learn\n(SaaS Product)", "#dae8fc", "#6c8ebf", "z4"))
    cells.append(rect("c3", 770, 250, 120, 60, "AG ONE Work", "#eeeeee", "#999999", "z4", 0))
    cells.append(rect("c4", 770, 320, 120, 60, "AG ONE Safe", "#eeeeee", "#999999", "z4", 0))

    cells.append(swimlane("z5", 340, 400, 780, 120, "Data & Events", "#e1d5e7", "#9673a6", "z2"))
    cells.append(cylinder("d1", 360, 440, 140, 60, "Azure SQL\nagone-prod", parent="z5"))
    cells.append(rect("d2", 530, 440, 140, 60, "Service Bus\nmarketplace-events", parent="z5"))
    cells.append(rect("d3", 700, 440, 120, 60, "Key Vault", parent="z5"))
    cells.append(rect("d4", 840, 440, 120, 60, "App Insights", parent="z5"))

    cells.append(rect("f1", 560, 560, 220, 50, "Fulfillment API v2", "#f8cecc", "#b85450", "z2"))

    cells += [
        edge("a1", "u1", "u2", "Subscribe"),
        edge("a2", "u2", "c1", "Redirect + token"),
        edge("a3", "c1", "f1", "Resolve / Activate"),
        edge("a4", "c1", "d1", "Provision tenant"),
        edge("a5", "c1", "d2", "Publish events"),
        edge("a6", "d2", "c2", "Async provisioning", dashed=True),
        edge("a7", "u1", "c1", "Login"),
        edge("a8", "c1", "u3", "OIDC / JWT"),
        edge("a9", "c1", "c2", "JWT cookie"),
        edge("a10", "c2", "d1", "learn.* data"),
        edge("a11", "c3", "c1", "SSO only", dashed=True),
        edge("a12", "c4", "c1", "SSO only", dashed=True),
    ]
    return cells


def page2_sequence():
    actors = [
        ("s1", 30, "Purchaser"), ("s2", 180, "Azure Marketplace"),
        ("s3", 350, "AG ONE Landing"), ("s4", 520, "Fulfillment API"),
        ("s5", 690, "AG ONE Gateway"), ("s6", 860, "Azure SQL"),
        ("s7", 1030, "AG ONE Learn"),
    ]
    cells = []
    y0 = 40
    for aid, x, name in actors:
        cells.append(rect(aid, x, y0, 120, 40, name, "#dae8fc", "#6c8ebf"))
        cells.append(
            f'<mxCell id="{aid}L" value="" style="endArrow=none;dashed=1;html=1;" edge="1" parent="1">'
            f'<mxGeometry relative="1" as="geometry">'
            f'<mxPoint x="{x+60}" y="90" as="sourcePoint"/>'
            f'<mxPoint x="{x+60}" y="700" as="targetPoint"/></mxGeometry></mxCell>'
        )
    steps = [
        (120, "s1", "s2", "Subscribe"),
        (150, "s2", "s3", "Redirect + token"),
        (180, "s3", "s4", "POST /resolve"),
        (210, "s4", "s3", "subscriptionId, planId"),
        (240, "s3", "s5", "Upsert MarketplaceSub"),
        (270, "s5", "s4", "POST /activate"),
        (300, "s5", "s6", "Create Tenant + Subscription"),
        (330, "s5", "s1", "Redirect to login"),
        (380, "s1", "s5", "Login"),
        (410, "s5", "s1", "JWT cookie"),
        (450, "s1", "s7", "Launch Learn"),
        (480, "s7", "s5", "Validate JWT"),
        (510, "s7", "s6", "Query learn.*"),
    ]
    for i, (y, src, tgt, lbl) in enumerate(steps):
        cells.append(edge(f"sq{i}", src, tgt, lbl))
    return cells


def page3_onboarding():
    cells = []
    flow = [
        (340, 40, "Azure Marketplace Purchase"),
        (340, 100, "Redirect to Landing Page\n?token=..."),
        (340, 170, "Token present?"),
        (120, 240, "Show error"),
        (560, 240, "Resolve token\nFulfillment API"),
        (560, 310, "Valid subscription?"),
        (560, 380, "Store MarketplaceSubscriptions\nPendingActivation"),
        (560, 450, "Activate subscription"),
        (560, 520, "Enqueue ProvisioningJob"),
        (560, 590, "Create core.Tenants"),
        (560, 660, "Create core.Subscriptions"),
        (340, 730, "Create Admin User + Roles"),
        (340, 800, "Seed learn.* defaults"),
        (340, 870, "Welcome email"),
        (340, 940, "Redirect to AG ONE Login"),
    ]
    ids = []
    for i, (x, y, lbl) in enumerate(flow):
        cid = f"o{i}"
        ids.append(cid)
        if "?" in lbl:
            cells.append(diamond(cid, x, y, 160, 70, lbl))
        elif "error" in lbl.lower():
            cells.append(rect(cid, x, y, 140, 50, lbl, "#f8cecc", "#b85450"))
        else:
            cells.append(rect(cid, x, y, 220, 50, lbl))
    for i in range(len(ids) - 1):
        cells.append(edge(f"oe{i}", ids[i], ids[i + 1]))
    cells.append(edge("oe_err", "o2", "o3", "No"))
    cells.append(edge("oe_yes", "o2", "o5", "Yes"))
    return cells


def page4_er():
    entities = [
        ("t1", 40, 40, 200, 120, "core.Tenants\n─────────\nId PK\nName, Slug\nProvisioningSource\nStatus"),
        ("t2", 300, 40, 200, 120, "core.Products\n─────────\nId PK\nCode\nMarketplaceOfferId\nIsMarketplaceEnabled"),
        ("t3", 560, 40, 220, 140, "core.Subscriptions\n─────────\nId PK\nTenantId FK\nProductId FK\nStatus, PlanName\nMaxLicenses\nBillingSource"),
        ("t4", 300, 240, 260, 160, "core.MarketplaceSubscriptions\n─────────\nId PK\nAzureSubscriptionId UK\nSubscriptionId FK\nOfferId, PlanId\nSaasStatus\nPurchaserEmail\nQuantity"),
        ("t5", 40, 240, 200, 100, "core.Users\n─────────\nId PK\nTenantId FK\nEmail"),
        ("t6", 40, 400, 200, 100, "core.UserRoles\n─────────\nUserId FK\nRoleId FK\nTenantId, ProductId"),
        ("t7", 300, 480, 200, 80, "core.MarketplaceWebhookEvents\n─────────\nIdempotencyKey UK\nAction, PayloadJson"),
        ("t8", 560, 480, 220, 80, "core.TenantProvisioningJobs\n─────────\nStatus, Step\nAttemptCount"),
    ]
    cells = [rect(*e) for e in entities]
    cells += [
        edge("r1", "t1", "t3", "1:N"),
        edge("r2", "t2", "t3", "1:N"),
        edge("r3", "t3", "t4", "1:1"),
        edge("r4", "t1", "t5", "1:N"),
        edge("r5", "t5", "t6", "1:N"),
        edge("r6", "t4", "t7", "1:N"),
        edge("r7", "t4", "t8", "1:N"),
    ]
    return cells


def page5_integration():
    cells = []
    cells.append(swimlane("ig", 40, 40, 900, 500, "AG ONE Gateway — Marketplace Integration"))
    boxes = [
        ("i1", 80, 100, 180, 60, "MarketplaceLanding\nController"),
        ("i2", 300, 100, 180, 60, "MarketplaceWebhook\nController"),
        ("i3", 520, 100, 200, 60, "MarketplaceFulfillment\nService"),
        ("i4", 760, 100, 160, 60, "TenantProvisioning\nService"),
        ("i5", 200, 240, 200, 60, "WebhookProcessor\nWorker"),
        ("i6", 480, 240, 200, 60, "MarketplaceLifecycle\nService"),
        ("i7", 300, 380, 220, 60, "Fulfillment API v2\n(Microsoft)"),
        ("i8", 560, 380, 180, 60, "Azure SQL\ncore.* + learn.*"),
    ]
    for b in boxes:
        cells.append(rect(*b, parent="ig"))
    cells.append(cloud("amp", 80, 380, 160, 70, "Azure Marketplace", parent="ig"))
    cells += [
        edge("i_e1", "amp", "i1", "Redirect", parent="ig"),
        edge("i_e2", "amp", "i2", "Webhooks", parent="ig"),
        edge("i_e3", "i1", "i3", parent="ig"),
        edge("i_e4", "i2", "i5", parent="ig"),
        edge("i_e5", "i5", "i3", parent="ig"),
        edge("i_e6", "i5", "i6", parent="ig"),
        edge("i_e7", "i3", "i7", parent="ig"),
        edge("i_e8", "i4", "i8", parent="ig"),
        edge("i_e9", "i6", "i8", parent="ig"),
    ]
    return cells


def page6_auth():
    cells = []
    actors = [("au1", 60, "User"), ("au2", 280, "AG ONE Gateway"), ("au3", 520, "Azure SQL"), ("au4", 760, "AG ONE Learn")]
    for aid, x, name in actors:
        cells.append(rect(aid, x, 40, 160, 50, name))
    steps = [
        (130, "au1", "au2", "POST /auth/login"),
        (170, "au2", "au3", "Find user + check Tenant/Subscription status"),
        (210, "au2", "au2", "Build JWT (tenant_id, subscription_status, permissions)"),
        (250, "au2", "au1", "Set agone_sso_token cookie"),
        (310, "au1", "au4", "Request with cookie"),
        (350, "au4", "au2", "Validate / refresh token"),
        (390, "au2", "au4", "ClaimsPrincipal"),
        (430, "au4", "au4", "Enforce tenant_id on learn.* queries"),
    ]
    for i, (y, src, tgt, lbl) in enumerate(steps):
        cells.append(edge(f"au_e{i}", src, tgt, lbl))
    return cells


def page7_system_flow():
    steps = [
        "① Browse Azure Marketplace",
        "② Subscribe",
        "③ Azure: PendingFulfillmentStart",
        "④ Redirect to Landing Page",
        "⑤ Resolve token",
        "⑥ Activate subscription",
        "⑦ Provision Tenant",
        "⑧ Create Subscription (Learn)",
        "⑨ Map AzureSub → Tenant",
        "⑩ Create Admin User + Roles",
        "⑪ Seed learn.* data",
        "⑫ Login via AG ONE",
        "⑬ Issue JWT",
        "⑭ Launch Learn",
        "⑮ Validate SSO",
        "⑯ Access granted",
    ]
    cells = []
    for i, s in enumerate(steps):
        y = 40 + i * 55
        cells.append(rect(f"sf{i}", 300, y, 280, 45, s))
        if i > 0:
            cells.append(edge(f"sf_e{i}", f"sf{i-1}", f"sf{i}"))
    return cells


def page8_deployment():
    cells = []
    cells.append(swimlane("dp", 20, 20, 1100, 620, "Azure Resource Group: rg-agone-prod"))
    cells.append(cloud("dpu", 60, 70, 180, 70, "Customers + Marketplace", parent="dp"))
    cells.append(rect("dpdns", 280, 70, 200, 50, "Azure DNS\nagone / learn", parent="dp"))
    cells.append(rect("dpagw", 520, 70, 200, 50, "App Gateway + WAF", parent="dp"))
    cells.append(rect("dpgw", 280, 180, 220, 80, "App: agone-gateway\nLanding · Auth · Webhooks", "#3b82f6", "#1e3a5f", "dp"))
    cells.append(rect("dplp", 540, 180, 220, 80, "App: agone-learn\nProduct API + UI", parent="dp"))
    cells.append(cylinder("dpsql", 280, 320, 200, 80, "Azure SQL\nagone-prod", parent="dp"))
    cells.append(rect("dpsb", 520, 320, 180, 80, "Service Bus", parent="dp"))
    cells.append(rect("dpkv", 730, 320, 140, 80, "Key Vault", parent="dp"))
    cells.append(rect("dpai", 900, 320, 160, 80, "App Insights", parent="dp"))
    cells.append(swimlane("dpaks", 280, 450, 560, 120, "Optional: AKS Scale-Out", "#eeeeee", "#999999", "dp"))
    cells.append(rect("dping", 300, 490, 160, 50, "Ingress Controller", parent="dpaks"))
    cells.append(rect("dppod1", 490, 490, 150, 50, "learn-api pods", parent="dpaks"))
    cells.append(rect("dppod2", 670, 490, 150, 50, "worker pods", parent="dpaks"))
    cells += [
        edge("dp1", "dpu", "dpdns", parent="dp"),
        edge("dp2", "dpdns", "dpagw", parent="dp"),
        edge("dp3", "dpagw", "dpgw", parent="dp"),
        edge("dp4", "dpagw", "dplp", parent="dp"),
        edge("dp5", "dpgw", "dpsql", parent="dp"),
        edge("dp6", "dpgw", "dpsb", parent="dp"),
        edge("dp7", "dplp", "dpsql", parent="dp"),
        edge("dp8", "dpgw", "dpkv", parent="dp"),
    ]
    return cells


def page9_lifecycle():
    states = [
        ("st0", 340, 40, 220, 50, "PendingFulfillmentStart"),
        ("st1", 340, 140, 220, 50, "Subscribed"),
        ("st2", 80, 240, 180, 50, "Suspended"),
        ("st3", 600, 240, 180, 50, "Unsubscribed"),
    ]
    cells = [rect(*s, fill="#fff2cc", stroke="#d6b656") for s in states]
    cells += [
        edge("ls1", "st0", "st1", "Activate (your API)"),
        edge("ls2", "st1", "st2", "Suspend webhook"),
        edge("ls3", "st2", "st1", "Reinstate webhook"),
        edge("ls4", "st1", "st3", "Unsubscribe webhook"),
        edge("ls5", "st0", "st3", "Timeout / cancel"),
    ]
    actions = [
        (40, 380, "Subscribe → Landing page resolve + activate"),
        (40, 420, "ChangePlan → Update PlanName + PATCH API"),
        (40, 460, "ChangeQuantity → Update MaxLicenses"),
        (40, 500, "Suspend → Block login, Status=Suspended"),
        (40, 540, "Reinstate → Restore Active"),
        (40, 580, "Unsubscribe → Cancel, soft-disable access"),
    ]
    for i, (x, y, lbl) in enumerate(actions):
        cells.append(rect(f"la{i}", x, y, 400, 35, lbl, "#dae8fc", "#6c8ebf"))
    return cells


def build_drawio():
    pages = [
        ("1-High-Level-Architecture", page1_high_level),
        ("2-End-to-End-Sequence", page2_sequence),
        ("3-SaaS-Onboarding-Flow", page3_onboarding),
        ("4-Database-ER-Diagram", page4_er),
        ("5-Integration-Design", page5_integration),
        ("6-Authentication-Flow", page6_auth),
        ("7-System-Flow-Steps", page7_system_flow),
        ("8-Deployment-Architecture", page8_deployment),
        ("9-Subscription-Lifecycle", page9_lifecycle),
    ]

    diagram_xml = []
    for name, builder in pages:
        cells = builder()
        body = "\n        ".join(cells)
        did = str(uuid.uuid4())
        diagram_xml.append(f"""  <diagram id="{did}" name="{esc(name)}">
    <mxGraphModel dx="1422" dy="794" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="1600" pageHeight="1200" math="0" shadow="0">
      <root>
        <mxCell id="0"/>
        <mxCell id="1" parent="0"/>
        {body}
      </root>
    </mxGraphModel>
  </diagram>""")

    ts = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%S.000Z")
    content = f"""<mxfile host="app.diagrams.net" modified="{ts}" agent="AG-ONE-Doc-Generator" version="22.1.0" type="device">
{chr(10).join(diagram_xml)}
</mxfile>"""
    DRAWIO_PATH.write_text(content, encoding="utf-8")
    print(f"Created: {DRAWIO_PATH}")


if __name__ == "__main__":
    build_drawio()
    build_word_document()
