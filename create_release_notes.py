#!/usr/bin/env python3
"""
AG ONE — Release Notes & History Document
Professional Word document covering all 5 products with full release timeline.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import date

# ─── Palette ──────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
BLUE   = RGBColor(0x3B, 0x82, 0xF6)
GREEN  = RGBColor(0x10, 0xB9, 0x81)
AMBER  = RGBColor(0xF5, 0x9E, 0x0B)
RED    = RGBColor(0xEF, 0x44, 0x44)
PURPLE = RGBColor(0x8B, 0x5C, 0xF6)
PINK   = RGBColor(0xEC, 0x48, 0x99)
DARK   = RGBColor(0x37, 0x41, 0x51)
GRAY   = RGBColor(0x6B, 0x72, 0x80)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

PRODUCT_COLORS = {
    "AG ONE":  NAVY,
    "OneWork": BLUE,
    "Learn":   PURPLE,
    "Safe":    GREEN,
    "OneHire": AMBER,
}


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
    paragraph._p.get_or_add_pPr().append(shading)


def add_styled_paragraph(doc, text, style="Normal", bold=False, color=DARK,
                          size=10, space_after=4, space_before=0, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = color
    run.font.size = Pt(size)
    run.font.name = "Aptos"
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align:
        p.alignment = align
    return p


def add_section_divider(doc, hex_color="3B82F6"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    set_paragraph_shading(p, hex_color)
    run = p.add_run(" ")
    run.font.size = Pt(2)


def add_release_entry(doc, version, release_date, release_type, status,
                       summary, features=None, fixes=None, hotfix_items=None,
                       known_issues=None, notes=None, product_color=NAVY):
    type_colors = {
        "MVP Release": "10B981",
        "Hotfix": "EF4444",
        "Minor Release": "3B82F6",
        "Major Release": "1E3A5F",
        "Patch": "F59E0B",
    }
    type_hex = type_colors.get(release_type, "3B82F6")

    # Version header bar
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = True

    # Version cell
    c0 = tbl.cell(0, 0)
    set_cell_shading(c0, "1E3A5F")
    p0 = c0.paragraphs[0]
    r0 = p0.add_run(f"  {version}")
    r0.bold = True
    r0.font.size = Pt(13)
    r0.font.color.rgb = WHITE
    r0.font.name = "Aptos"
    c0.width = Inches(2.8)

    # Type cell
    c1 = tbl.cell(0, 1)
    set_cell_shading(c1, type_hex)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(f" {release_type} ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = WHITE
    r1.font.name = "Aptos"
    c1.width = Inches(1.8)

    # Date + status cell
    c2 = tbl.cell(0, 2)
    set_cell_shading(c2, "F3F4F6")
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(f"{release_date}  •  {status}  ")
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK
    r2.font.name = "Aptos"
    c2.width = Inches(2.4)

    # Remove table borders
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tcBorders>'
            )
            tcPr.append(tcBorders)

    # Summary
    add_styled_paragraph(doc, summary, bold=False, color=DARK, size=10,
                          space_before=6, space_after=4)

    # Features
    if features:
        add_styled_paragraph(doc, "New Features & Enhancements", bold=True,
                              color=product_color, size=10, space_before=4, space_after=2)
        for f in features:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f)
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK
            run.font.name = "Aptos"
            p.paragraph_format.space_after = Pt(1)

    # Bug Fixes
    if fixes:
        add_styled_paragraph(doc, "Bug Fixes", bold=True, color=product_color,
                              size=10, space_before=4, space_after=2)
        for f in fixes:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f)
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK
            run.font.name = "Aptos"
            p.paragraph_format.space_after = Pt(1)

    # Hotfix items
    if hotfix_items:
        add_styled_paragraph(doc, "Hotfix Details", bold=True, color=RED,
                              size=10, space_before=4, space_after=2)
        for f in hotfix_items:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f)
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK
            run.font.name = "Aptos"
            p.paragraph_format.space_after = Pt(1)

    # Known issues
    if known_issues:
        add_styled_paragraph(doc, "Known Issues", bold=True, color=AMBER,
                              size=10, space_before=4, space_after=2)
        for f in known_issues:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f)
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK
            run.font.name = "Aptos"
            p.paragraph_format.space_after = Pt(1)

    # Notes
    if notes:
        add_styled_paragraph(doc, "Notes", bold=True, color=GRAY,
                              size=10, space_before=4, space_after=2)
        add_styled_paragraph(doc, notes, color=GRAY, size=9.5, space_after=2)

    # Spacer
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_product_section(doc, product_name, product_color, releases, upcoming=None):
    """Full product section: title banner, timeline summary, release entries, upcoming."""
    hex_color = str(product_color)

    # Product banner
    p = doc.add_paragraph()
    set_paragraph_shading(p, hex_color)
    run = p.add_run(f"  {product_name}  —  Release History")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = WHITE
    run.font.name = "Aptos"
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)

    # Release timeline summary table
    tbl = doc.add_table(rows=1 + len(releases), cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    headers = ["Version", "Date", "Type", "Status", "Summary"]
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        set_cell_shading(cell, hex_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = WHITE
        r.font.name = "Aptos"

    for ri, rel in enumerate(releases, 1):
        vals = [rel["version"], rel["date"], rel["type"], rel["status"], rel["summary_short"]]
        for ci, val in enumerate(vals):
            cell = tbl.cell(ri, ci)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci < 4 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.size = Pt(9)
            r.font.color.rgb = DARK
            r.font.name = "Aptos"
            if ri % 2 == 0:
                set_cell_shading(cell, "F3F4F6")

    # Set column widths
    for row in tbl.rows:
        row.cells[0].width = Inches(1.2)
        row.cells[1].width = Inches(1.2)
        row.cells[2].width = Inches(1.2)
        row.cells[3].width = Inches(1.0)
        row.cells[4].width = Inches(2.6)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Detailed release entries
    add_styled_paragraph(doc, "Detailed Release Notes", bold=True, color=product_color,
                          size=14, space_before=8, space_after=6)

    for rel in releases:
        add_release_entry(
            doc,
            version=rel["version"],
            release_date=rel["date"],
            release_type=rel["type"],
            status=rel["status"],
            summary=rel["summary"],
            features=rel.get("features"),
            fixes=rel.get("fixes"),
            hotfix_items=rel.get("hotfix_items"),
            known_issues=rel.get("known_issues"),
            notes=rel.get("notes"),
            product_color=product_color,
        )

    # Upcoming releases placeholder
    if upcoming:
        add_styled_paragraph(doc, "Upcoming Releases", bold=True, color=product_color,
                              size=14, space_before=12, space_after=6)
        tbl2 = doc.add_table(rows=1 + len(upcoming), cols=4)
        tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl2.style = "Table Grid"
        for i, h in enumerate(["Version", "Target Date", "Type", "Planned Scope"]):
            cell = tbl2.cell(0, i)
            set_cell_shading(cell, hex_color)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h)
            r.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = WHITE
            r.font.name = "Aptos"
        for ri, up in enumerate(upcoming, 1):
            for ci, val in enumerate(up):
                cell = tbl2.cell(ri, ci)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci < 3 else WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(val)
                r.font.size = Pt(9)
                r.font.color.rgb = DARK
                r.font.name = "Aptos"
                if ri % 2 == 0:
                    set_cell_shading(cell, "F3F4F6")
        for row in tbl2.rows:
            row.cells[0].width = Inches(1.2)
            row.cells[1].width = Inches(1.4)
            row.cells[2].width = Inches(1.2)
            row.cells[3].width = Inches(3.4)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════
doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ─── COVER PAGE ───────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
set_paragraph_shading(p, "1E3A5F")
run = p.add_run("\n  AG ONE Platform\n")
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = WHITE
run.font.name = "Aptos"
p.paragraph_format.space_after = Pt(0)

p2 = doc.add_paragraph()
set_paragraph_shading(p2, "3B82F6")
run2 = p2.add_run("\n  Release Notes & Version History\n")
run2.bold = True
run2.font.size = Pt(20)
run2.font.color.rgb = WHITE
run2.font.name = "Aptos"
p2.paragraph_format.space_after = Pt(0)

p3 = doc.add_paragraph()
set_paragraph_shading(p3, "DBEAFE")
run3 = p3.add_run(f"\n  2026  •  All Products  •  Confidential\n")
run3.font.size = Pt(12)
run3.font.color.rgb = NAVY
run3.font.name = "Aptos"

doc.add_paragraph()
doc.add_paragraph()

# Products list
add_styled_paragraph(doc, "Products Covered in This Document:", bold=True,
                      color=NAVY, size=14, space_before=12)

products_list = [
    ("AG ONE", "Identity & Access Management Platform — Panel 1"),
    ("OneWork", "Workforce Management Platform"),
    ("Learn", "Learning Management System (LMS)"),
    ("Safe", "Compliance & Risk Management"),
    ("OneHire", "Recruitment & Applicant Tracking"),
]
for pname, pdesc in products_list:
    p = doc.add_paragraph()
    run = p.add_run(f"    {pname}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = PRODUCT_COLORS[pname]
    run.font.name = "Aptos"
    run2 = p.add_run(f"  —  {pdesc}")
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY
    run2.font.name = "Aptos"
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
add_styled_paragraph(doc, "Document Information", bold=True, color=NAVY,
                      size=12, space_before=8)

info_table = doc.add_table(rows=5, cols=2)
info_table.style = "Table Grid"
info_data = [
    ("Document Owner", "Engineering Manager"),
    ("Last Updated", date.today().strftime("%d %B %Y")),
    ("Version", "1.0"),
    ("Classification", "Internal — Confidential"),
    ("Distribution", "Directors, Tech Leads, Engineering Team"),
]
for i, (k, v) in enumerate(info_data):
    c0 = info_table.cell(i, 0)
    set_cell_shading(c0, "1E3A5F")
    p = c0.paragraphs[0]
    r = p.add_run(f"  {k}")
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = WHITE
    r.font.name = "Aptos"
    c0.width = Inches(2)

    c1 = info_table.cell(i, 1)
    p = c1.paragraphs[0]
    r = p.add_run(f"  {v}")
    r.font.size = Pt(9.5)
    r.font.color.rgb = DARK
    r.font.name = "Aptos"

doc.add_page_break()

# ─── TABLE OF CONTENTS PAGE ──────────────────────────────────────────
add_styled_paragraph(doc, "Table of Contents", bold=True, color=NAVY,
                      size=18, space_before=4, space_after=12)

toc_items = [
    "1.  AG ONE (Panel 1)  —  Release History",
    "        v1.0.0  MVP Release  •  v1.0.1  Hotfix 1  •  v1.0.2  Hotfix 2",
    "        Upcoming: v1.1.0, v1.2.0, v2.0.0",
    "",
    "2.  OneWork  —  Release History",
    "        v1.0.0  MVP Release",
    "        Upcoming: v1.1.0, v1.2.0",
    "",
    "3.  Learn  —  Release History",
    "        v1.0.0  MVP Release",
    "        Upcoming: v1.1.0, v1.2.0",
    "",
    "4.  Safe  —  Release History",
    "        v1.0.0  MVP Release",
    "        Upcoming: v1.1.0, v1.2.0",
    "",
    "5.  OneHire  —  Release History",
    "        (No releases yet — planned)",
    "",
    "6.  Release Type Legend & Versioning Policy",
]
for item in toc_items:
    if not item:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        continue
    is_sub = item.startswith("    ")
    p = add_styled_paragraph(
        doc, item,
        bold=not is_sub,
        color=NAVY if not is_sub else GRAY,
        size=11 if not is_sub else 9.5,
        space_after=2,
    )

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════
# 1. AG ONE — PANEL 1 (fully populated: MVP + 2 hotfixes)
# ═══════════════════════════════════════════════════════════════════════
ag_one_releases = [
    {
        "version": "v1.0.0",
        "date": "29 Mar 2026",
        "type": "MVP Release",
        "status": "Live",
        "summary_short": "MVP launch — full IAM platform",
        "summary": "Initial production release of the AG ONE platform. Includes complete Identity & Access Management, multi-tenant architecture, SSO integration, role-based permissions, and admin panel.",
        "features": [
            "User Management — full CRUD with search, filters, pagination, and bulk operations",
            "Role Management — create/edit roles with granular permission assignment, system roles (read-only), and role deduplication",
            "Permission Matrix — hierarchical permission system across all products (AgOne, Work, Learn, Safe, Pulse, Spot)",
            "Tenant Management — company profile, SSO settings (Azure AD), API key generation & revocation",
            "Assign Access Wizard — 3-step wizard: select users → choose products & roles → review & confirm",
            "Audit Trail — full audit logging with search, pagination, 7-day summary stats (user changes, role changes, API calls)",
            "Login History — track user logins with IP, browser, OS, location, success/failure status",
            "SSO Middleware — automatic Bearer token forwarding on all HttpClient calls via TokenForwardingHandler",
            "Master Data Tenant — tenant list view with toggle controls, filter dropdowns, and edit capability",
            "My Products — product card layout with launch buttons for subscribed products",
            "Responsive UI — Blazor WASM with Inter font, blue/navy design system, 8–12px rounded corners",
        ],
        "fixes": [
            "Fixed checkbox visibility: increased disabled opacity, preserved blue accent on checked+disabled state",
            "Fixed Roles.razor: resolved missing closing div for rw-footer causing cascading layout errors",
            "Fixed pre-check: added dedicated user-roles endpoint, matched ProductId null→Guid.Empty conversion",
            "Fixed AssignAccess: load products+roles upfront, fixed Guid.Empty vs null matching for platform roles",
            "Fixed products not showing: built allProducts from permission matrix API, used nullable Guid for product selection",
            "Fixed Roles deduplication: deduplicated API rows, fixed empty-state for Guid.Empty product, fixed auto-check on edit",
        ],
        "known_issues": [
            "Pagination: max 5 page numbers shown with ellipsis — works correctly but may need UX review for very large datasets",
        ],
        "notes": "Sprint 1–6 delivery. All features built with full team (Geena, Nastaran). Launched on schedule with zero P0 bugs.",
    },
    {
        "version": "v1.0.1",
        "date": "02 Apr 2026",
        "type": "Hotfix",
        "status": "Live",
        "summary_short": "Critical hotfix — SSO token & role sync",
        "summary": "Emergency hotfix addressing SSO token refresh failures and role synchronisation issues reported by production tenants after MVP launch.",
        "hotfix_items": [
            "Fixed SSO token refresh loop — TokenForwardingHandler was not correctly detecting expired tokens, causing infinite redirect loops for some tenants",
            "Fixed role sync delay — UserRoles cache (10-min TTL) was not invalidated after Assign Access wizard completion, causing stale permissions for up to 10 minutes",
            "Fixed tenant switching — switching between tenants in multi-tenant mode did not clear the cached permission set, showing stale permissions from the previous tenant",
        ],
        "fixes": [
            "Added force-cache-clear on Assign Access completion to ensure immediate permission reflection",
            "Added tenant_id validation on every permission cache lookup to prevent cross-tenant cache hits",
            "Improved SSO middleware error handling with proper 401 response instead of redirect loop",
        ],
        "notes": "Deployed within 48 hours of MVP launch. Affected 3 tenants. No data loss. Root cause: cache key did not include tenant context.",
    },
    {
        "version": "v1.0.2",
        "date": "08 Apr 2026",
        "type": "Hotfix",
        "status": "Live",
        "summary_short": "Hotfix — audit log & API key fixes",
        "summary": "Second hotfix addressing audit log display issues and API key creation edge case reported during Sprint 7 stabilisation.",
        "hotfix_items": [
            "Fixed audit log pagination — page count was calculated incorrectly when filters were applied, showing blank pages at the end",
            "Fixed API key creation — duplicate key names were allowed, causing confusion in the API Keys management panel",
            "Fixed login history timezone — timestamps were displayed in UTC instead of tenant's configured timezone",
        ],
        "fixes": [
            "Audit log: recalculated total count after filter application, not before",
            "API keys: added unique name validation before creation with clear error message",
            "Login history: applied tenant timezone offset to all displayed timestamps",
            "Minor UI fix: API key revoke confirmation dialog was not closing on successful revocation",
        ],
        "notes": "Deployed during Sprint 7. Low severity — no tenant impact on core functionality. Improved overall admin panel reliability.",
    },
]

ag_one_upcoming = [
    ("v1.1.0", "May 2026", "Minor Release", "Advanced user analytics, bulk role operations, enhanced audit dashboard with charts"),
    ("v1.2.0", "Jul 2026", "Minor Release", "Multi-product admin console, cross-product role mapping, API marketplace v1"),
    ("v2.0.0", "Oct 2026", "Major Release", "Enterprise SSO (SAML + OIDC), advanced RBAC policies, white-label tenant branding, performance at scale"),
]

add_product_section(doc, "AG ONE  —  Panel 1", NAVY, ag_one_releases, ag_one_upcoming)


# ═══════════════════════════════════════════════════════════════════════
# 2. OneWork (MVP done)
# ═══════════════════════════════════════════════════════════════════════
onework_releases = [
    {
        "version": "v1.0.0",
        "date": "29 Mar 2026",
        "type": "MVP Release",
        "status": "Live",
        "summary_short": "MVP launch — workforce management",
        "summary": "Initial production release of the OneWork platform. Full employee management, dynamic workflow engine, task management, and manager dashboards.",
        "features": [
            "Employee Management — full CRUD with search, filters, team hierarchy views",
            "Workflow Engine — dynamic workflow builder with task assignment and notifications",
            "Task Management — create, assign, track, and close tasks with status workflows",
            "Manager Dashboard — team views, KPI widgets, real-time metrics",
            "Reporting Module — export-ready reports, filterable data views",
        ],
        "fixes": [
            "42 bugs fixed during Sprint 5 integration testing",
            "Zero critical bugs in final pre-launch testing round",
        ],
        "notes": "Sprint 1–6 delivery. Full team: Abdullah, Nastaran, Jawad, Geena. Launched on schedule.",
    },
]

onework_upcoming = [
    ("v1.1.0", "Jun 2026", "Minor Release", "Advanced workflow templates, email notifications, calendar integration"),
    ("v1.2.0", "Aug 2026", "Minor Release", "Reporting v2 with charts, mobile-responsive views, third-party integrations"),
    ("v2.0.0", "Nov 2026", "Major Release", "AI-powered insights, advanced automation rules, enterprise-scale performance"),
]

add_product_section(doc, "OneWork", BLUE, onework_releases, onework_upcoming)


# ═══════════════════════════════════════════════════════════════════════
# 3. Learn (MVP done)
# ═══════════════════════════════════════════════════════════════════════
learn_releases = [
    {
        "version": "v1.0.0",
        "date": "29 Mar 2026",
        "type": "MVP Release",
        "status": "Live",
        "summary_short": "MVP launch — LMS platform",
        "summary": "Initial production release of Learn LMS. Courses, learning paths, assessments, certifications, and progress tracking.",
        "features": [
            "Course Management — create courses with drag-and-drop content builder, multi-format uploads",
            "Learning Paths — structured learning journeys with sequential/parallel modules",
            "Assessment Engine — quiz builder with multiple question types, auto-grading",
            "Certifications — auto-generated certificates on course/path completion",
            "Progress Tracking — learner dashboards with completion %, time spent, scores",
            "Auto-Assignment — rule-based course assignment by role, department, or custom criteria",
        ],
        "fixes": [
            "All acceptance criteria passed in Sprint 5 testing",
        ],
        "notes": "Sprint 1–6 delivery. Team: Ricky, Loc, Than. Smooth launch with zero post-launch incidents.",
    },
]

learn_upcoming = [
    ("v1.1.0", "Jun 2026", "Minor Release", "Advanced analytics dashboard, learner recommendations, completion reports"),
    ("v1.2.0", "Aug 2026", "Minor Release", "Content marketplace, SCORM support, mobile-responsive learning"),
    ("v2.0.0", "Nov 2026", "Major Release", "AI-powered recommendations, adaptive learning paths, full mobile app"),
]

add_product_section(doc, "Learn", PURPLE, learn_releases, learn_upcoming)


# ═══════════════════════════════════════════════════════════════════════
# 4. Safe (MVP done)
# ═══════════════════════════════════════════════════════════════════════
safe_releases = [
    {
        "version": "v1.0.0",
        "date": "29 Mar 2026",
        "type": "MVP Release",
        "status": "Live",
        "summary_short": "MVP launch — compliance platform",
        "summary": "Initial production release of Safe compliance platform. Policy management, compliance tracking, audit workflows, risk matrix, and reporting dashboards.",
        "features": [
            "Policy Management — create, version, and publish compliance policies with approval workflows",
            "Data Library — centralised document store for compliance evidence and artefacts",
            "Compliance Tracking — checklist-based compliance monitoring with automated reminders",
            "Audit Workflows — structured audit trails with findings, actions, and sign-offs",
            "Risk Matrix — visual risk assessment with likelihood/impact scoring",
            "Compliance Dashboard — real-time compliance posture with export-ready reports",
        ],
        "fixes": [
            "Passed security audit with zero critical findings in Sprint 5",
        ],
        "notes": "Sprint 1–6 delivery. Team: Faisal, Surya, Kiritini. Launched on schedule after passing security audit.",
    },
]

safe_upcoming = [
    ("v1.1.0", "Jun 2026", "Minor Release", "Regulatory update automation, scheduled compliance checks, advanced reporting"),
    ("v1.2.0", "Sep 2026", "Minor Release", "Third-party integrations (GRC tools), automated alerts, evidence collection"),
    ("v2.0.0", "Dec 2026", "Major Release", "Enterprise compliance suite with AI-driven risk predictions"),
]

add_product_section(doc, "Safe", GREEN, safe_releases, safe_upcoming)


# ═══════════════════════════════════════════════════════════════════════
# 5. OneHire (no releases yet)
# ═══════════════════════════════════════════════════════════════════════
add_styled_paragraph(doc, "", size=2)
p = doc.add_paragraph()
set_paragraph_shading(p, "F59E0B")
run = p.add_run("  OneHire  —  Release History")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = WHITE
run.font.name = "Aptos"
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(8)

add_styled_paragraph(doc, "No releases yet. Development planned to begin in Q2 2026.",
                      color=GRAY, size=11, space_before=8, space_after=4)
add_styled_paragraph(doc, "Planned Releases", bold=True, color=AMBER,
                      size=14, space_before=12, space_after=6)

onehire_upcoming = [
    ("v0.1.0", "May 2026", "Alpha", "Core job posting and applicant tracking features"),
    ("v0.5.0", "Jun 2026", "Beta", "Interview scheduling, candidate pipeline, basic reporting"),
    ("v1.0.0", "Jul 2026", "MVP Release", "Full recruitment portal with applicant tracking and analytics"),
    ("v1.1.0", "Sep 2026", "Minor Release", "Advanced hiring analytics, integration with HR systems"),
]

tbl_oh = doc.add_table(rows=1 + len(onehire_upcoming), cols=4)
tbl_oh.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl_oh.style = "Table Grid"
for i, h in enumerate(["Version", "Target Date", "Type", "Planned Scope"]):
    cell = tbl_oh.cell(0, i)
    set_cell_shading(cell, "F59E0B")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = WHITE
    r.font.name = "Aptos"
for ri, up in enumerate(onehire_upcoming, 1):
    for ci, val in enumerate(up):
        cell = tbl_oh.cell(ri, ci)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci < 3 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(val)
        r.font.size = Pt(9)
        r.font.color.rgb = DARK
        r.font.name = "Aptos"
        if ri % 2 == 0:
            set_cell_shading(cell, "F3F4F6")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════
# RELEASE TYPE LEGEND & VERSIONING POLICY
# ═══════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
set_paragraph_shading(p, "1E3A5F")
run = p.add_run("  Release Type Legend & Versioning Policy")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = WHITE
run.font.name = "Aptos"
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(12)

# Version format
add_styled_paragraph(doc, "Versioning Format:  vMAJOR.MINOR.PATCH", bold=True,
                      color=NAVY, size=12, space_after=8)

version_rules = [
    ("MAJOR (vX.0.0)", "Breaking changes, major architecture changes, or significant new feature sets"),
    ("MINOR (v1.X.0)", "New features, enhancements, and non-breaking improvements"),
    ("PATCH (v1.0.X)", "Bug fixes, hotfixes, and minor corrections"),
]
for label, desc in version_rules:
    p = doc.add_paragraph()
    run = p.add_run(f"    {label}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    run.font.name = "Aptos"
    run2 = p.add_run(f"  —  {desc}")
    run2.font.size = Pt(10)
    run2.font.color.rgb = DARK
    run2.font.name = "Aptos"
    p.paragraph_format.space_after = Pt(3)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Release types
add_styled_paragraph(doc, "Release Types", bold=True, color=NAVY, size=12, space_after=6)

types_table = doc.add_table(rows=6, cols=3)
types_table.style = "Table Grid"
types_data = [
    ("Type", "Description", "Approval Required"),
    ("MVP Release", "First production release of a product", "Director + Tech Lead"),
    ("Major Release", "Significant new features or breaking changes", "Director + Tech Lead"),
    ("Minor Release", "New features & enhancements, backward compatible", "Tech Lead"),
    ("Hotfix", "Critical production fix, expedited deployment", "Tech Lead (post-mortem required)"),
    ("Patch", "Non-critical bug fixes, scheduled deployment", "Tech Lead"),
]
type_hex_map = {
    "Type": "1E3A5F",
    "MVP Release": "10B981",
    "Major Release": "1E3A5F",
    "Minor Release": "3B82F6",
    "Hotfix": "EF4444",
    "Patch": "F59E0B",
}
for ri, row_data in enumerate(types_data):
    for ci, val in enumerate(row_data):
        cell = types_table.cell(ri, ci)
        p = cell.paragraphs[0]
        r = p.add_run(f"  {val}")
        r.font.name = "Aptos"
        if ri == 0:
            set_cell_shading(cell, "1E3A5F")
            r.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = WHITE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            r.font.size = Pt(9.5)
            r.font.color.rgb = DARK
            if ci == 0:
                hex_c = type_hex_map.get(val, "3B82F6")
                set_cell_shading(cell, hex_c)
                r.bold = True
                r.font.color.rgb = WHITE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ─── SAVE ─────────────────────────────────────────────────────────────
output_path = "/workspace/AG_ONE_Release_Notes_2026.docx"
doc.save(output_path)
print(f"✅ Saved: {output_path}")
